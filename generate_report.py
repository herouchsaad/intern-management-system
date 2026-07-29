import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading("État d'avancement du PFA", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Mise en place d'une plateforme DevOps de gestion de stage")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("Ce document résume l'état d'avancement actuel du Projet de Fin d'Études (PFA) concernant la mise en place d'une plateforme DevOps pour l'application de gestion de stage (Intern Management System).")
    
    doc.add_heading("2. Réalisations actuelles", level=1)
    
    doc.add_heading("Découverte et analyse de l'application Web :", level=2)
    doc.add_paragraph("L'application a été identifiée et analysée. Elle se compose de trois parties principales :", style='List Bullet')
    doc.add_paragraph("Un frontend développé en React.", style='List Bullet 2')
    doc.add_paragraph("Un backend développé en Node.js (Express) utilisant TypeScript.", style='List Bullet 2')
    doc.add_paragraph("Une base de données relationnelle PostgreSQL.", style='List Bullet 2')
    
    doc.add_heading("Conteneurisation (Dockerisation) :", level=2)
    doc.add_paragraph("Une architecture basée sur des conteneurs a été mise en place avec succès pour standardiser les environnements de développement et de production. L'approche DevOps commence par la création de ces artefacts de déploiement isolés :")
    
    doc.add_paragraph("Frontend : Création d'un Dockerfile multi-étapes (multi-stage build). La première étape compile l'application React avec Node.js, et la seconde utilise un serveur web léger Nginx (nginx:1.25-alpine) pour servir les fichiers statiques de l'application sur le port 80. Une configuration Nginx personnalisée a été intégrée.", style='List Bullet')
    
    doc.add_paragraph("Backend : Création d'un Dockerfile multi-étapes. La première étape gère l'installation des dépendances complètes et la compilation du code TypeScript en JavaScript, tandis que la seconde exécute le serveur Node.js allégé (sans les dépendances de développement) sur le port 5000 pour minimiser la surface d'attaque et la taille de l'image.", style='List Bullet')
    
    doc.add_paragraph("Base de données : Utilisation de l'image officielle postgres:15-alpine. La persistance des données est assurée via des volumes Docker (postgres_data). Un script SQL d'initialisation (ims-tables.sql) est monté pour créer automatiquement le schéma lors du premier démarrage.", style='List Bullet')
    
    doc.add_paragraph("Orchestration (Docker Compose) : Création d'un fichier docker-compose.yml à la racine du projet permettant de lancer et de lier les 3 services simultanément avec la commande docker-compose up --build. Un réseau interne (ims_network) a été configuré pour permettre une communication sécurisée entre les conteneurs (par exemple, le backend et Nginx interagissent avec ce réseau privé). Des 'healthchecks' (vérifications de santé) ont été ajoutés pour s'assurer que les services démarrent dans le bon ordre.", style='List Bullet')
    
    doc.add_heading("3. Prochaines étapes (Perspectives DevOps)", level=1)
    doc.add_paragraph("Mise en place des pipelines CI/CD (Intégration et Déploiement Continus) via GitHub Actions (ou GitLab CI) pour automatiser les tests, le build des images Docker, et le déploiement.", style='List Bullet')
    doc.add_paragraph("Déploiement de l'application sur un cluster Kubernetes (K8s) ou sur une infrastructure Cloud pour assurer la scalabilité et la haute disponibilité.", style='List Bullet')
    doc.add_paragraph("Intégration d'outils de monitoring et de logging (comme Prometheus, Grafana, ou la suite ELK) pour la supervision des conteneurs en production.", style='List Bullet')
    
    doc.save("Etat_Avancement_PFA.docx")
    print("Le document Word a été généré avec succès sous le nom 'Etat_Avancement_PFA.docx'")

if __name__ == '__main__':
    create_report()
